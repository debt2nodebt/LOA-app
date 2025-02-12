import streamlit as st
import pandas as pd
from docx import Document
import os

# Define relative path for the Word template
TEMPLATE_PATH = os.path.join("template", "LOA_template.docx")


# Function to round off values to the nearest hundred
def round_off(value):
    return round(value / 100) * 100

# Function to modify the Word template
from docx.shared import Pt

def modify_word_template(client_data, loan_details, client_name):
    doc = Document(TEMPLATE_PATH)

    # Replace placeholders outside the table
    for para in doc.paragraphs:
        for key, value in client_data.items():
            if key in para.text:
                para.text = para.text.replace(key, str(value))
                for run in para.runs:
                    run.font.size = Pt(12)  # Set font size to 12

    # Initialize total calculations
    total_balance_os = 0
    total_25 = 0
    total_30 = 0

    # Modify tables
    for table in doc.tables:
        # First Table (Contains {clientname}, {dateofbirth}, {address})
        if "{clientname}" in table.rows[0].cells[0].text:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in client_data.items():
                        if key in cell.text:
                            cell.text = cell.text.replace(key, str(value))
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    run.font.size = Pt(12)

        # Second Table (Creditor Details)
        elif "Name of Creditor (App Loan/Bank name)" in table.rows[0].cells[0].text:
            for i, loan in enumerate(loan_details):
                if i > 0:
                    table.add_row()

                row = table.rows[i + 1].cells
                row[0].text = loan["Name of Creditor (App Loan/Bank name)"]
                row[1].text = loan["Type of Debt/ Loan"]
                row[2].text = loan["Loan Account Number"]

                balance_os = float(loan["Balance O/S"])
                approx_25 = round_off(balance_os * 0.25)
                approx_30 = round_off(balance_os * 0.30)

                row[3].text = str(balance_os)  
                row[4].text = str(approx_25)  
                row[5].text = str(approx_30)  

                # Add values to totals
                total_balance_os += balance_os
                total_25 += approx_25
                total_30 += approx_30

                # Set font size for each cell
                for cell in row:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(12)

            # Add Total Row
            total_row = table.add_row().cells
            total_row[0].text = "Total"
            total_row[1].text = ""
            total_row[2].text = ""
            total_row[3].text = str(total_balance_os)
            total_row[4].text = str(total_25)
            total_row[5].text = str(total_30)

            # Apply font size for total row
            for cell in total_row:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(12)
                        run.bold = True  # Make totals bold

    # Save file as "LOA of {client_name}.docx"
    file_name = f"LOA of {client_name}.docx"
    doc.save(file_name)
    return file_name


# Streamlit UI
st.title("LOA Generator")

# Client Details Input
st.subheader("Client Details")
client_name = st.text_input("Client Name")
client_city = st.text_input("Client City")
client_dob = st.date_input("Date of Birth")
client_address = st.text_area("Address")
date = st.date_input("Date")

# Loan Details Input
st.subheader("Loan Details")
loan_data = pd.DataFrame(columns=[
    "Name of Creditor (App Loan/Bank name)", 
    "Type of Debt/ Loan", 
    "Loan Account Number", 
    "Balance O/S"
])

loan_data = st.data_editor(loan_data, num_rows="dynamic")

# Generate LOA Button
if st.button("Generate LOA"):
    if not client_name or loan_data.empty:
        st.error("Please fill in all required fields.")
    else:
        client_info = {
            "{clientname}": client_name,
            "{clientcity}": client_city,
            "{date}": date.strftime("%d-%m-%Y"),
            "{dateofbirth}": client_dob.strftime("%d-%m-%Y"),
            "{address}": client_address
        }

        # Convert DataFrame to list of dictionaries
        loan_details = loan_data.to_dict(orient="records")

        # Modify the Word template
        output_file = modify_word_template(client_info, loan_details, client_name)

        # Display Download Button After Generation
        with open(output_file, "rb") as file:
            st.download_button(
                label="Download LOA",
                data=file,
                file_name=output_file,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
